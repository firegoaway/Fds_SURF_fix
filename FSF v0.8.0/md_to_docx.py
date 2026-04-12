"""
Модуль конвертации Markdown → DOCX с поддержкой:
- Таблицы с форматированием
- Математические формулы (LaTeX → OMML)
- Изображения (внедрение в DOCX)
- Заголовки, списки, цитаты, код
- Сохранение стилей и табуляции

Использует Pandoc (приоритет) или python-docx (резерв).
"""

import os
import re
import base64
import subprocess
import tempfile
import shutil
from pathlib import Path
from typing import Optional, List, Tuple

# Путь к директории с модулем для доступа к шаблону и pandoc
MODULE_DIR = Path(__file__).parent.resolve()
TEMPLATE_DOCX = MODULE_DIR / "tu_aupt_report2.docx"
PANDOC_EXE = MODULE_DIR / "pandoc_embed" / "pandoc.exe"

# Попытка импорта python-docx для резервного варианта
try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# Попытка импорта markdown для парсинга
try:
    import markdown
    from markdown.extensions.tables import TableExtension
    from markdown.extensions.fenced_code import FencedCodeExtension
    HAS_MARKDOWN = True
except ImportError:
    HAS_MARKDOWN = False

# Попытка импорта BeautifulSoup для парсинга HTML
try:
    from bs4 import BeautifulSoup
    HAS_BEAUTIFULSOUP = True
except ImportError:
    HAS_BEAUTIFULSOUP = False


class MarkdownToDocxConverter:
    """
    Мощный конвертер Markdown в DOCX с поддержкой формул, таблиц и изображений.
    
    Поддерживаемые элементы:
    - Заголовки (H1-H6)
    - Параграфы, списки (маркированные/нумерованные)
    - Таблицы с форматированием
    - Математические формулы LaTeX ($$...$$ и $...$)
    - Изображения (PNG, JPG, GIF, BMP)
    - Код (блоки и инлайн)
    - Цитаты, горизонтальные линии
    - Жирный, курсив, подчёркнутый текст
    """
    
    def __init__(self, use_pandoc: bool = True, pandoc_path: Optional[str] = None):
        """
        Инициализация конвертера.
        
        Args:
            use_pandoc: Использовать Pandoc (True) или python-docx (False)
            pandoc_path: Путь к исполняемому файлу Pandoc (если не в PATH)
        """
        self.use_pandoc = use_pandoc
        self.pandoc_path = pandoc_path or self._find_pandoc()
        self.pandoc_available = self.pandoc_path is not None and self._check_pandoc()
        
        # Если Pandoc недоступен, пробуем резервный вариант
        if self.use_pandoc and not self.pandoc_available:
            print("⚠️ Pandoc не найден. Переключение на резервный режим (python-docx).")
            print("   Для полной поддержки формул установите Pandoc: https://pandoc.org/installing.html")
            self.use_pandoc = False
        
        if not self.use_pandoc and not (HAS_DOCX and HAS_MARKDOWN):
            raise ImportError(
                "Необходимые библиотеки не установлены.\n"
                "Установите: pip install python-docx markdown beautifulsoup4\n"
                "Для полной поддержки формул также установите Pandoc."
            )
    
    def _find_pandoc(self) -> Optional[str]:
        """Поиск Pandoc: сначала локальная папка pandoc_embed, затем системный PATH."""
        # 1. Проверяем локальный pandoc_embed/pandoc.exe
        if PANDOC_EXE.is_file():
            return str(PANDOC_EXE)

        # 2. Проверка стандартных путей
        possible_paths = [
            "pandoc",
            r"C:\Program Files\Pandoc\pandoc.exe",
            r"C:\Program Files (x86)\Pandoc\pandoc.exe",
            os.path.expanduser(r"~\AppData\Local\Pandoc\pandoc.exe"),
        ]

        for path in possible_paths:
            if os.path.isfile(path) and os.access(path, os.X_OK):
                return path

        # 3. Поиск через where (Windows)
        try:
            result = subprocess.run(
                ["where", "pandoc"],
                capture_output=True,
                text=True,
                timeout=5
            )
            if result.returncode == 0:
                return result.stdout.strip().split('\n')[0].strip()
        except Exception:
            pass

        return None
    
    def _check_pandoc(self) -> bool:
        """Проверка работоспособности Pandoc."""
        try:
            result = subprocess.run(
                [self.pandoc_path, "--version"],
                capture_output=True,
                text=True,
                timeout=10
            )
            return result.returncode == 0
        except Exception:
            return False
    
    def convert(self, md_path: str, docx_path: str, 
                preserve_images: bool = True,
                image_folder: Optional[str] = None) -> bool:
        """
        Конвертация MD файла в DOCX.
        
        Args:
            md_path: Путь к исходному MD файлу
            docx_path: Путь для сохранения DOCX
            preserve_images: Сохранять изображения
            image_folder: Папка с изображениями (если None, ищутся рядом с MD)
        
        Returns:
            True при успехе
        """
        md_path = Path(md_path).resolve()
        docx_path = Path(docx_path).resolve()
        
        if not md_path.exists():
            raise FileNotFoundError(f"MD файл не найден: {md_path}")
        
        if self.use_pandoc and self.pandoc_available:
            return self._convert_with_pandoc(md_path, docx_path, preserve_images, image_folder)
        else:
            return self._convert_with_python_docx(md_path, docx_path, preserve_images, image_folder)
    
    def _convert_with_pandoc(self, md_path: Path, docx_path: Path,
                             preserve_images: bool, image_folder: Optional[str]) -> bool:
        """Конвертация через Pandoc (максимальное качество)."""

        # Подготовка аргументов для Pandoc
        # Pandoc 3.x использует UTF-8 по умолчанию
        args = [
            self.pandoc_path,
            str(md_path),
            "-f", "markdown+tex_math_dollars+tex_math_single_backslash",
            "-t", "docx",
            "-o", str(docx_path),
            "--reference-doc", str(TEMPLATE_DOCX),
            "--standalone",
            "--toc",  # Оглавление
        ]

        # Добавляем путь к изображениям
        if preserve_images and image_folder:
            args.extend(["--resource-path", image_folder])
        elif preserve_images:
            args.extend(["--resource-path", str(md_path.parent)])

        try:
            result = subprocess.run(
                args,
                capture_output=True,
                encoding='utf-8',
                timeout=120,
                cwd=str(md_path.parent)
            )

            if result.returncode != 0:
                print(f"⚠️ Pandoc вернул ошибку: {result.stderr}")
                print(f"   Пробую упрощённый режим...")
                return self._convert_with_pandoc_fallback(md_path, docx_path, preserve_images, image_folder)

            # Выводим предупреждения Pandoc (если есть)
            if result.stderr:
                print(f"⚠️ Предупреждения Pandoc: {result.stderr}")

            # Постобработка таблиц для единообразного форматирования
            print(f"📝 Постобработка таблиц...")
            self.postprocess_tables(str(docx_path))

            return True

        except subprocess.TimeoutExpired:
            print("⚠️ Превышено времени ожидания Pandoc. Пробую упрощённый режим...")
            return self._convert_with_pandoc_fallback(md_path, docx_path, preserve_images, image_folder)
        except Exception as e:
            import traceback
            print(f"⚠️ Ошибка Pandoc: {e}. Пробую резервный режим...")
            traceback.print_exc()
            return self._convert_with_pandoc_fallback(md_path, docx_path, preserve_images, image_folder)
    
    def _convert_with_pandoc_fallback(self, md_path: Path, docx_path: Path,
                                      preserve_images: bool, image_folder: Optional[str]) -> bool:
        """Упрощённая конвертация Pandoc (без некоторых опций)."""
        args = [
            self.pandoc_path,
            str(md_path),
            "-f", "markdown+tex_math_dollars",
            "-t", "docx",
            "-o", str(docx_path),
            "--reference-doc", str(TEMPLATE_DOCX),
            "--standalone",
        ]

        if preserve_images:
            res_path = image_folder if image_folder else str(md_path.parent)
            args.extend(["--resource-path", res_path])

        try:
            result = subprocess.run(
                args,
                capture_output=True,
                encoding='utf-8',
                timeout=120,
                cwd=str(md_path.parent)
            )
            # Постобработка таблиц для единообразного форматирования
            if result.returncode == 0:
                self.postprocess_tables(str(docx_path))
            return result.returncode == 0
        except Exception as e:
            import traceback
            print(f"⚠️ Ошибка упрощённого Pandoc: {e}")
            traceback.print_exc()
            return False
    
    def _convert_with_python_docx(self, md_path: Path, docx_path: Path,
                                  preserve_images: bool, image_folder: Optional[str]) -> bool:
        """
        Резервная конвертация через python-docx.
        Поддерживает базовое форматирование, таблицы, изображения.
        Формулы LaTeX конвертируются в текст (требует дополнительного обработчика).
        """
        if not (HAS_DOCX and HAS_MARKDOWN and HAS_BEAUTIFULSOUP):
            print("❌ Отсутствуют необходимые библиотеки для резервного режима.")
            return False
        
        # Чтение MD файла
        with open(md_path, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        # Обработка формул LaTeX перед конвертацией
        md_content = self._preprocess_latex_formulas(md_content)
        
        # Парсинг Markdown в HTML
        md_extensions = ['tables', 'fenced_code', 'codehilite', 'toc', 'nl2br']
        html = markdown.markdown(md_content, extensions=md_extensions)
        
        # Парсинг HTML
        soup = BeautifulSoup(html, 'html.parser')
        
        # Создание DOCX документа
        doc = Document()
        self._setup_docx_styles(doc)
        
        # Обработка изображений
        if image_folder is None:
            image_folder = str(md_path.parent)

        # Конвертация HTML элементов в DOCX
        self._process_html_elements(soup, doc, str(md_path.parent), preserve_images)

        # Сохранение документа
        doc.save(str(docx_path))
        
        # Постобработка таблиц для единообразного форматирования
        self.postprocess_tables(str(docx_path))
        
        return True
    
    def _preprocess_latex_formulas(self, md_content: str) -> str:
        """
        Предобработка формул LaTeX для совместимости.
        Конвертирует $$...$$ и $...$ в формат для последующей обработки.
        """
        # Защита формул от Markdown-парсера
        # Блочные формулы $$...$$
        md_content = re.sub(
            r'\$\$(.+?)\$\$',
            lambda m: f'`LATEX_BLOCK:{m.group(1).strip()}`',
            md_content,
            flags=re.DOTALL
        )
        
        # Инлайн формулы $...$
        md_content = re.sub(
            r'\$([^$]+?)\$',
            lambda m: f'`LATEX_INLINE:{m.group(1).strip()}`',
            md_content
        )
        
        return md_content
    
    def _setup_docx_styles(self, doc: Document):
        """Настройка стилей DOCX документа."""
        # Настройка стилей заголовков
        styles = doc.styles
        
        # Основной стиль
        if 'Normal' in styles:
            normal_style = styles['Normal']
            normal_font = normal_style.font
            normal_font.name = 'Times New Roman'
            normal_font.size = Pt(12)
        
        # Стили заголовков
        heading_styles = [
            ('Heading 1', 16, True),
            ('Heading 2', 14, True),
            ('Heading 3', 12, True),
            ('Heading 4', 12, False),
            ('Heading 5', 11, False),
            ('Heading 6', 10, False),
        ]
        
        for style_name, size, bold in heading_styles:
            if style_name in styles:
                style = styles[style_name]
                style.font.size = Pt(size)
                style.font.bold = bold
                style.font.name = 'Times New Roman'
    
    def _process_html_elements(self, soup: BeautifulSoup, doc: Document,
                               base_path: str, preserve_images: bool):
        """Рекурсивная обработка HTML элементов."""
        
        for element in soup.children:
            self._process_element(element, doc, base_path, preserve_images)
    
    def _process_element(self, element, doc: Document, base_path: str, 
                         preserve_images: bool, paragraph=None):
        """Обработка отдельного HTML элемента."""
        
        if element.name is None:  # Текстовый узел
            text = element.string.strip() if element.string else ''
            if text and paragraph is not None:
                run = paragraph.add_run(text)
                # Проверка на формулы LaTeX
                if text.startswith('LATEX_BLOCK:') or text.startswith('LATEX_INLINE:'):
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
            return
        
        tag_name = element.name.lower()
        
        # Заголовки
        if tag_name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            level = int(tag_name[1])
            p = doc.add_paragraph(style=f'Heading {level}')
            self._process_inline_elements(element, p)
        
        # Параграфы
        elif tag_name == 'p':
            p = doc.add_paragraph()
            self._process_inline_elements(element, p)
        
        # Списки
        elif tag_name == 'ul':
            for li in element.find_all('li', recursive=False):
                p = doc.add_paragraph(style='List Bullet')
                self._process_inline_elements(li, p)
        
        elif tag_name == 'ol':
            for i, li in enumerate(element.find_all('li', recursive=False), 1):
                p = doc.add_paragraph(style='List Number')
                self._process_inline_elements(li, p)
        
        # Таблицы
        elif tag_name == 'table':
            self._process_table(element, doc)
        
        # Изображения
        elif tag_name == 'img' and preserve_images:
            src = element.get('src', '')
            if src:
                self._add_image(doc, src, base_path)
        
        # Код
        elif tag_name == 'pre':
            code = element.get_text()
            p = doc.add_paragraph()
            run = p.add_run(code)
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            # Добавляем фон для блока кода
            shading_elm = OxmlElement('w:pShd')
            shading_elm.set(qn('w:val'), 'clear')
            shading_elm.set(qn('w:fill'), 'F5F5F5')
            p._element.append(shading_elm)
        
        # Цитаты
        elif tag_name == 'blockquote':
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.5)
            p.paragraph_format.right_indent = Cm(1.5)
            run = p.add_run(element.get_text())
            run.italic = True
        
        # Горизонтальная линия
        elif tag_name == 'hr':
            p = doc.add_paragraph()
            run = p.add_run('─' * 50)
            run.font.size = Pt(8)
        
        # Рекурсивная обработка вложенных элементов
        else:
            for child in element.children:
                self._process_element(child, doc, base_path, preserve_images, paragraph)
    
    def _process_inline_elements(self, element, paragraph):
        """Обработка инлайн элементов (bold, italic, code, etc.)."""
        
        for child in element.children:
            if isinstance(child, str):
                text = child.strip()
                if text:
                    paragraph.add_run(text + ' ')
            elif child.name:
                run = paragraph.add_run()
                
                if child.name == 'strong' or child.name == 'b':
                    run.add_text(child.get_text())
                    run.bold = True
                elif child.name == 'em' or child.name == 'i':
                    run.add_text(child.get_text())
                    run.italic = True
                elif child.name == 'u':
                    run.add_text(child.get_text())
                    run.underline = True
                elif child.name == 'code':
                    run.add_text(child.get_text())
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
                elif child.name == 'a':
                    run.add_text(child.get_text())
                    run.underline = True
                    run.font.color.rgb = RGBColor(0, 0, 255)
                else:
                    # Рекурсивная обработка
                    self._process_inline_elements(child, paragraph)
    
    def _process_table(self, table_element, doc: Document):
        """Обработка HTML таблицы."""
        
        rows = table_element.find_all('tr')
        if not rows:
            return
        
        # Определение размеров таблицы
        num_cols = 0
        for row in rows:
            cells = row.find_all(['td', 'th'])
            num_cols = max(num_cols, len(cells))
        
        if num_cols == 0:
            return
        
        # Создание таблицы
        table = doc.add_table(rows=len(rows), cols=num_cols)
        table.style = 'Table Grid'
        
        # Заполнение таблицы
        for i, row in enumerate(rows):
            cells = row.find_all(['td', 'th'])
            is_header = row.find('th') is not None
            
            for j, cell in enumerate(cells):
                if j < num_cols:
                    table_cell = table.cell(i, j)
                    p = table_cell.paragraphs[0]
                    p.clear()
                    
                    # Добавление текста ячейки
                    text = cell.get_text().strip()
                    run = p.add_run(text)
                    
                    if is_header:
                        run.bold = True
                        run.font.size = Pt(11)
                    
                    # Выравнивание по центру для заголовков
                    if is_header:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _add_image(self, doc: Document, src: str, base_path: str):
        """Добавление изображения в документ."""

        # Обработка относительных путей
        if not os.path.isabs(src):
            image_path = os.path.join(base_path, src)
        else:
            image_path = src

        # Проверка существования файла
        if os.path.exists(image_path):
            try:
                # Добавление изображения с центрированием
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(image_path, width=Inches(6))
            except Exception as e:
                print(f"⚠️ Не удалось добавить изображение {src}: {e}")
        else:
            print(f"⚠️ Изображение не найдено: {image_path}")

    def postprocess_tables(self, docx_path: str):
        """
        Постобработка таблиц в DOCX файле: применение автоподбора по содержимому.

        Для каждой таблицы:
        1. Устанавливается w:type=auto для tblW (автоподбор по содержимому)
        2. Удаляется tblLayout с fixed (чтобы Word сам определял ширину столбцов)
        3. Форматируются шрифты и границы ячеек

        Args:
            docx_path: Путь к DOCX файлу
        """
        if not HAS_DOCX:
            print("⚠️ python-docx не доступен, постобработка невозможна")
            return False

        try:
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement

            doc = Document(docx_path)
            tables_modified = 0

            for table in doc.tables:
                tables_modified += 1

                # === АВТОПОДБОР ПО СОДЕРЖИМОМУ ===
                tblPr = table._element.tblPr
                if tblPr is None:
                    tblPr = OxmlElement('w:tblPr')
                    table._element.insert(0, tblPr)

                # Находим и удаляем tblW (фиксированную ширину)
                tblW = tblPr.find(qn('w:tblW'))
                if tblW is not None:
                    tblPr.remove(tblW)

                # Создаём tblW с типом "auto" — автоподбор по содержимому
                tblW_new = OxmlElement('w:tblW')
                tblW_new.set(qn('w:w'), '0')
                tblW_new.set(qn('w:type'), 'auto')
                tblPr.insert(0, tblW_new)

                # Удаляем tblLayout с fixed — заменяем на autofit
                tblLayout = tblPr.find(qn('w:tblLayout'))
                if tblLayout is not None:
                    tblPr.remove(tblLayout)
                tblLayout_new = OxmlElement('w:tblLayout')
                tblLayout_new.set(qn('w:type'), 'autofit')
                tblPr.append(tblLayout_new)

                # НЕ удаляем tblGrid/gridCol — пусть Word сам подберет ширину

                # Форматирование ячеек
                for row in table.rows:
                    for cell in row.cells:
                        # Устанавливаем вертикальное выравнивание по центру
                        for paragraph in cell.paragraphs:
                            paragraph_format = paragraph.paragraph_format
                            paragraph_format.line_spacing = Pt(14)

                        # Устанавливаем границы ячеек
                        self._set_cell_borders(cell)

            if tables_modified > 0:
                doc.save(docx_path)
                print(f"✅ Отформатировано таблиц: {tables_modified}")
            else:
                print("ℹ️ Таблицы не найдены")

            return True

        except Exception as e:
            print(f"⚠️ Ошибка постобработки таблиц: {e}")
            import traceback
            traceback.print_exc()
            return False

    def _set_cell_borders(self, cell):
        """
        Установка границ для ячейки таблицы.
        
        Args:
            cell: Ячейка таблицы python-docx
        """
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        
        # Создаём элемент границ
        tcBorders = OxmlElement('w:tcBorders')
        
        # Стили границ
        border_style = 'single'
        border_sz = '4'  # 4 * 1/8 pt = 0.5 pt
        border_color = '000000'
        
        # Создаём границы для всех сторон
        for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{side}')
            border.set(qn('w:val'), border_style)
            border.set(qn('w:sz'), border_sz)
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), border_color)
            tcBorders.append(border)
        
        tcPr.append(tcBorders)

    def convert_string(self, md_content: str, docx_path: str,
                       preserve_images: bool = True,
                       image_folder: Optional[str] = None) -> bool:
        """
        Конвертация MD строки в DOCX.
        
        Args:
            md_content: Содержимое MD в виде строки
            docx_path: Путь для сохранения DOCX
            preserve_images: Сохранять изображения
            image_folder: Папка с изображениями
        
        Returns:
            True при успехе
        """
        # Создание временного MD файла
        with tempfile.NamedTemporaryFile(mode='w', suffix='.md', 
                                         delete=False, encoding='utf-8') as f:
            f.write(md_content)
            temp_md = f.name
        
        try:
            result = self.convert(temp_md, docx_path, preserve_images, image_folder)
            return result
        finally:
            # Очистка временного файла
            if os.path.exists(temp_md):
                os.unlink(temp_md)


def convert_md_to_docx(md_path: str, docx_path: str, 
                       use_pandoc: bool = True,
                       preserve_images: bool = True,
                       image_folder: Optional[str] = None) -> bool:
    """
    Удобная функция для конвертации MD → DOCX.
    
    Args:
        md_path: Путь к MD файлу
        docx_path: Путь для DOCX файла
        use_pandoc: Использовать Pandoc (рекомендуется)
        preserve_images: Сохранять изображения
        image_folder: Папка с изображениями
    
    Returns:
        True при успехе
    
    Example:
        >>> convert_md_to_docx('report.md', 'report.docx')
        True
    """
    converter = MarkdownToDocxConverter(use_pandoc=use_pandoc)
    return converter.convert(md_path, docx_path, preserve_images, image_folder)


# ============================================================================
# АВТОРЫ
# Дипова Н.Г.
# ----------------------------------------------------------------------------
#                               FIREGOAWAY NPO.
# ============================================================================

if __name__ == '__main__':
    # Пример использования
    import sys
    
    if len(sys.argv) < 3:
        print("Использование: python md_to_docx.py <input.md> <output.docx>")
        print("\nПример:")
        print("  python md_to_docx.py report.md report.docx")
        sys.exit(1)
    
    md_file = sys.argv[1]
    docx_file = sys.argv[2]
    
    print(f"Конвертация {md_file} → {docx_file}...")
    
    success = convert_md_to_docx(md_file, docx_file)
    
    if success:
        print(f"✅ Успешно! Файл сохранён: {docx_file}")
    else:
        print("❌ Ошибка конвертации. Проверьте логи выше.")
        sys.exit(1)
